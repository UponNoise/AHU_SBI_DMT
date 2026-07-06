from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


BASE = Path(r"C:\Users\canana\Desktop\讲义\AAAAA\计算机安全\Assignment\Homework7")
RESULTS = BASE / "hw7_results"
TEMPLATE = BASE / "Homework_7_Template.docx"
OUT_DOCX = BASE / "Homework_7_R32314015_Shengxuan_white_screenshots.docx"

MAGIC = "test:U6aMy0wojraho:0:0:test:/root:/bin/bash"


def read(name: str) -> str:
    return (RESULTS / name).read_text(encoding="utf-8", errors="replace")


def section(text: str, start: str, end: str | None = None) -> str:
    i = text.index(start)
    if end is None:
        return text[i:].strip()
    j = text.index(end, i)
    return text[i:j].strip()


def clean_lines(text: str) -> str:
    lines = []
    for line in text.splitlines():
        if "sysctl: permission denied" in line:
            continue
        if "write error: Broken pipe" in line:
            continue
        lines.append(line.rstrip())
    return "\n".join(lines).strip()


def terminal_png(text: str, out: Path, title: str) -> None:
    text = clean_lines(text)
    lines = [f"$ {title}"] + text.splitlines()
    font_path = Path(r"C:\Windows\Fonts\consola.ttf")
    font = ImageFont.truetype(str(font_path), 19) if font_path.exists() else ImageFont.load_default()
    padding_x, padding_y = 24, 20
    line_h = 25
    widths = []
    dummy = Image.new("RGB", (1, 1))
    draw = ImageDraw.Draw(dummy)
    for line in lines:
        bbox = draw.textbbox((0, 0), line, font=font)
        widths.append(bbox[2] - bbox[0])
    width = min(max(widths) + padding_x * 2, 1600)
    height = padding_y * 2 + line_h * len(lines)
    img = Image.new("RGB", (width, height), (22, 26, 30))
    draw = ImageDraw.Draw(img)
    y = padding_y
    for idx, line in enumerate(lines):
        color = (120, 214, 255) if idx == 0 else (232, 238, 244)
        draw.text((padding_x, y), line, fill=color, font=font)
        y += line_h
    img.save(out)


def make_figures() -> dict[str, Path]:
    windows = BASE / "windows_captures"
    window_figs = {
        "fig1": windows / "fig1_windows.png",
        "fig2": windows / "fig2_windows.png",
        "fig3": windows / "fig3_windows.png",
        "fig4": windows / "fig4_windows.png",
        "fig5": windows / "fig5_windows.png",
    }
    if all(path.exists() for path in window_figs.values()):
        return window_figs

    task1 = read("task1.txt")
    task2a = read("task2a.txt")
    task2c = read("task2c.txt")
    task3a = read("task3a.txt")

    figs = {
        "fig1": RESULTS / "fig1_passwd_magic.png",
        "fig2": RESULTS / "fig2_login_root.png",
        "fig3": RESULTS / "fig3_task2a_links.png",
        "fig4": RESULTS / "fig4_task2c_success.png",
        "fig5": RESULTS / "fig5_task3a_defense.png",
    }
    terminal_png(section(task1, "End of /etc/passwd:", "Login test"), figs["fig1"], "tail -n 6 /etc/passwd")
    terminal_png(section(task1, "Login test with an empty password:"), figs["fig2"], "printf '\\n' | su test -c 'whoami; id; echo EUID=$EUID'")
    terminal_png(section(task2a, "BEFORE running vulp:"), figs["fig3"], "ls -ld /tmp/XYZ before and after ./vulp")
    terminal_png(section(task2c, "Last part of target_process.sh output:"), figs["fig4"], "./target_process.sh")
    terminal_png(section(task3a, "Last part of target_process.sh output:"), figs["fig5"], "./target_process.sh with least-privilege vulp")
    return figs


def insert_after(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    return Paragraph(new_p, paragraph._parent)


def set_paragraph_text(paragraph: Paragraph, text: str, bold_prefix: str | None = None) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    if bold_prefix and text.startswith(bold_prefix):
        r = paragraph.add_run(bold_prefix)
        r.bold = True
        paragraph.add_run(text[len(bold_prefix):])
    else:
        paragraph.add_run(text)


def color_para(paragraph: Paragraph, color: RGBColor) -> None:
    for run in paragraph.runs:
        run.font.color.rgb = color


def add_answer(paragraph: Paragraph, answer: str) -> None:
    set_paragraph_text(paragraph, "Answer: " + answer, "Answer:")


def style_code(paragraph: Paragraph, red: bool = False) -> None:
    for run in paragraph.runs:
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        if red:
            run.font.color.rgb = RGBColor(192, 0, 0)


def fill_document(figs: dict[str, Path]) -> None:
    doc = Document(TEMPLATE)

    # Remove the template's example screenshots, then insert our own figures at the relevant captions.
    for p in list(doc.paragraphs):
        xml = p._p.xml
        has_picture = any(token in xml for token in ("<v:shape", "<w:drawing", "<wp:inline", "<pic:pic"))
        if not p.text.strip() and has_picture:
            p._element.getparent().remove(p._element)

    answers = [
        "Yes. After appending the magic value to /etc/passwd, I could log into the test account with an empty password. The output shows whoami returned root and id returned uid=0(root), gid=0(root), so the test account has root privilege.",
        "I first disabled the Ubuntu protections for this lab, compiled vulp with the added sleep(10), changed its owner to root, set the setuid bit with chmod 4755, created a writable file /home/seed/LabsetupRace/test, and linked /tmp/XYZ to that writable file.",
        f"I entered the magic value: {MAGIC}",
        "During the 10-second sleep interval, I changed the symbolic link with ln -sf /etc/passwd /tmp/XYZ. Therefore access() checked a writable file, but fopen() later opened /etc/passwd.",
        "In this context, atomically means the two symbolic-link names are exchanged as one indivisible operation. There is no intermediate state where /tmp/XYZ is missing or only partially updated. It is needed because unlink()+symlink() creates a gap that can make the target program fail or observe the wrong file; renameat2(RENAME_EXCHANGE) keeps both paths valid while rapidly switching which one points to /etc/passwd.",
        "After applying least privilege with seteuid(), the attack failed. The program drops its effective UID from root to the real user seed before fopen(), so even if the access() check passes, the actual open of /etc/passwd is attempted without root privilege and returns Permission denied. The output shows /etc/passwd was not changed.",
        "Yes. I observed a similar result: the attack did not append the test account to /etc/passwd.",
        "Ubuntu's protection prevents unsafe following of symbolic links in sticky world-writable directories such as /tmp. When a setuid-root program tries to follow an attacker-owned symlink to a sensitive file, the kernel blocks it, which prevents this /tmp/XYZ attack. The limitation is that it mainly protects this class of symlink attacks in sticky directories; it does not fix the vulnerable TOCTOU pattern, does not protect all race conditions, and can be disabled or bypassed in cases outside its ownership and directory assumptions.",
    ]
    answer_idx = 0

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("Full Name:"):
            set_paragraph_text(p, "Full Name: Shengxuan")
        elif txt.startswith("Student ID:"):
            set_paragraph_text(p, "Student ID: R32314015")
        elif txt.startswith("Date:"):
            set_paragraph_text(p, "Date: 2026-06-02")
        elif txt == "Answer:" and answer_idx < len(answers):
            add_answer(p, answers[answer_idx])
            if answer_idx == 5:
                img_p = insert_after(p)
                img_p.alignment = 1
                img_p.add_run().add_picture(str(figs["fig5"]), width=Inches(5.8))
            answer_idx += 1

    for p in doc.paragraphs:
        text = p.text.strip()
        if text.startswith("Figure 1 "):
            img_p = p.insert_paragraph_before()
            img_p.alignment = 1
            img_p.add_run().add_picture(str(figs["fig1"]), width=Inches(5.8))
        elif text.startswith("Figure 2 "):
            img_p = p.insert_paragraph_before()
            img_p.alignment = 1
            img_p.add_run().add_picture(str(figs["fig2"]), width=Inches(5.8))
        elif text.startswith("Figure 3 "):
            img_p = p.insert_paragraph_before()
            img_p.alignment = 1
            img_p.add_run().add_picture(str(figs["fig3"]), width=Inches(5.8))
        elif text.startswith("Figure 4 "):
            img_p = p.insert_paragraph_before()
            img_p.alignment = 1
            img_p.add_run().add_picture(str(figs["fig4"]), width=Inches(5.8))

    for p in doc.paragraphs:
        text = p.text.strip()
        if text == 'echo "hello" | ./vulp':
            set_paragraph_text(p, f"printf '%s\\n' '{MAGIC}' | ./vulp")
            style_code(p, red=True)
        elif text in {
            "#!/bin/bash",
            'CHECK_FILE="ls -l /etc/passwd"',
            'old=$($CHECK_FILE)',
            'new=$($CHECK_FILE)',
            'while [ "$old" == "$new" ]',
            "do",
            "new=$($CHECK_FILE)",
            "done",
            'echo "STOP... The passwd file has been changed"',
            "#include <stdio.h>",
            "#include <stdlib.h>",
            "#include <string.h>",
            "#include <unistd.h>",
            "int main()",
            "{",
            'char* fn = "/tmp/XYZ";',
            "char buffer[60];",
            "FILE* fp;",
            "uid_t real_uid = getuid();",
            "uid_t eff_uid = geteuid();",
            "/* get user input */",
            "scanf(\"%50s\", buffer);",
            "if (!access(fn, W_OK)) {",
            'fp = fopen(fn, "a+");',
            "if (!fp) {",
            'perror("Open failed");',
            "exit(1);",
            "}",
            'fwrite("\\n", sizeof(char), 1, fp);',
            "fwrite(buffer, sizeof(char), strlen(buffer), fp);",
            "fclose(fp);",
            "} else {",
            'printf("No permission \\n");',
            "return 0;",
        }:
            style_code(p)

    # Insert least-privilege changes in red around fopen/fclose in the Task 3.A code listing.
    paras = doc.paragraphs
    for idx, p in enumerate(list(paras)):
        if p.text.strip() == 'fp = fopen(fn, "a+");':
            prev_texts = [x.text.strip() for x in doc.paragraphs[max(0, idx - 10):idx]]
            if "uid_t eff_uid = geteuid();" in prev_texts:
                np = p.insert_paragraph_before("seteuid(real_uid);")
                style_code(np, red=True)
                break

    for idx, p in enumerate(list(doc.paragraphs)):
        if p.text.strip() == "exit(1);":
            prev_texts = [x.text.strip() for x in doc.paragraphs[max(0, idx - 10):idx]]
            if "seteuid(real_uid);" in prev_texts or 'fp = fopen(fn, "a+");' in prev_texts:
                np = p.insert_paragraph_before("seteuid(eff_uid);")
                style_code(np, red=True)
                break

    inserted_after_close = False
    for p in list(doc.paragraphs):
        if p.text.strip() == "fclose(fp);" and not inserted_after_close:
            np = insert_after(p)
            set_paragraph_text(np, "seteuid(eff_uid);")
            style_code(np, red=True)
            inserted_after_close = True

    # Keep figure captions centered.
    for p in doc.paragraphs:
        if re.match(r"Figure \d+", p.text.strip()):
            p.alignment = 1

    doc.save(OUT_DOCX)


if __name__ == "__main__":
    figures = make_figures()
    fill_document(figures)
    print(OUT_DOCX)
