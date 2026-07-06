# -*- coding: utf-8 -*-
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "doc"
MATERIALS_JSON = DOC_DIR / "materials_extract.json"
OUT = DOC_DIR / "Computer_Security_Final_Review_CN_EN.docx"


def set_run_font(run, size=10.5, bold=False, color=None):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_para(doc, text="", style=None, size=10.5, bold=False, color=None, space_after=4):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    return p


def add_bullet(doc, cn, en=None, level=0):
    text = cn if en is None else f"{cn}\nEN: {en}"
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    set_run_font(r, 10)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_number(doc, cn, en=None):
    text = cn if en is None else f"{cn}\nEN: {en}"
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_run_font(r, 10)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size={1: 16, 2: 13, 3: 11}.get(level, 11), bold=True, color=(31, 78, 121))
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_shading(hdr[i], "D9EAF7")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, 9.5, bold=True)
        if widths:
            hdr[i].width = widths[i]
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            set_run_font(r, 9)
            p.paragraph_format.space_after = Pt(0)
            if widths:
                cells[i].width = widths[i]
    doc.add_paragraph()
    return table


def extract_answer_keys(data):
    rows = []
    for name, text in data.get("practice", {}).items():
        part_i = re.search(r"Part I Solution:\s*([^P]+?)(?:Part II|$)", text, re.S)
        part_ii = re.search(r"Part II Solution:\s*([^P]+?)(?:Part III|$)", text, re.S)
        rows.append([
            name,
            re.sub(r"\s+", " ", part_i.group(1)).strip() if part_i else "See source",
            re.sub(r"\s+", " ", part_ii.group(1)).strip() if part_ii else "See source",
        ])
    for name, text in data.get("quiz", {}).items():
        qs = re.findall(r"(\d+)\.\s*\(\s*([A-DTF])\s*\)", text)
        rows.append([
            name,
            "; ".join(f"{n}:{ans}" for n, ans in qs) or "Short/fill-in answers embedded in source",
            "Short-answer reasoning must be memorized, not only the letter.",
        ])
    return rows


def main():
    data = json.loads(MATERIALS_JSON.read_text(encoding="utf-8"))

    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.65)
    sec.bottom_margin = Inches(0.65)
    sec.left_margin = Inches(0.72)
    sec.right_margin = Inches(0.72)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Computer Security Fundamentals Final Review\n计算机安全基础期末复习资料")
    set_run_font(r, 18, bold=True, color=(31, 78, 121))
    add_para(doc, "Evidence base: local Teams/course materials already present under 计算机安全: Practice, Quiz, PPT, Syllabus, Homework. 证据基础：本地已存在的 Teams/课程资料。", size=9.5)
    add_para(doc, "Priority rule: Practice and Quiz are treated as exam-style evidence; PPT and homework topics are used to fill concepts and examples. 优先规则：Practice/Quiz 优先，课件和作业用于补充知识点。", size=9.5)

    add_h(doc, "1. 范围结论 / Scope Conclusion", 1)
    add_bullet(doc, "高优先级：密码学、认证/PKI/TLS、访问控制、浏览器同源策略、CSRF/XSS/SQL 注入、DNS/BGP、DoS、Worms、IDS、防火墙。", "High priority: crypto, authentication/PKI/TLS, access control, SOP, CSRF/XSS/SQL injection, DNS/BGP, DoS, worms, IDS, firewalls.")
    add_bullet(doc, "中优先级：Shellcode、Sandbox、Memory Errors/Buffer Overflow、Race Condition。它们在作业和 PPT 中出现，但本地 Practice/Quiz 覆盖不足。", "Medium priority: shellcode, sandbox, memory errors/buffer overflow, race condition. They appear in homework/PPT but have weaker local practice/quiz evidence.")
    add_bullet(doc, "资料缺口：本地没有 Quiz 4-6；PPT 只到 CSE331-20 Memory Errors。Syllabus 后半部分还列出 ROP/Heap、Kernel、Program Analysis、Symbolic Execution、Side Channels、Ethics、HW/ML/IoT Security，若教师明天考全范围，这些需要至少扫标题和定义。", "Gap: local files do not include Quiz 4-6 and slides stop at CSE331-20. The syllabus lists later topics; skim definitions if the final is cumulative/full-syllabus.")
    add_bullet(doc, "实际备考策略：先刷 Practice/Quiz 到能解释原因，再用 PPT 补齐概念，最后只扫低证据主题。", "Strategy: redo practice/quiz until you can explain the reasoning, then use slides to patch concepts, then skim lower-evidence topics.")

    add_h(doc, "2. 明天考试前复习顺序 / Last-Day Order", 1)
    for cn, en in [
        ("90 分钟：重做 6 个 Quiz。不要只背答案，必须能说明为什么错项错误。", "90 min: redo all 6 quizzes. Explain why wrong choices are wrong."),
        ("150 分钟：刷 3 套 Practice，优先短答和计算题，因为它们最接近分析题。", "150 min: redo the 3 practice sets, prioritizing short answers and calculations."),
        ("90 分钟：按下面主题表背定义、攻击流程、防御方法。", "90 min: memorize definitions, attack flows, and defenses using the topic guide below."),
        ("45 分钟：系统安全补漏：shellcode、sandbox、memory errors、race condition。", "45 min: patch system-security gaps: shellcode, sandbox, memory errors, race condition."),
        ("30 分钟：只看答案索引和术语表，停止扩展新资料。", "30 min: review answer keys and glossary only; stop expanding scope."),
    ]:
        add_number(doc, cn, en)

    add_h(doc, "3. 高频知识点 / High-Yield Topics", 1)
    topics = [
        ("密码学 / Cryptography", [
            ("Kerckhoff 原则：算法公开仍应安全；不要依赖 Security through Obscurity。", "Kerckhoff's principle: the system remains secure even if the algorithm is public."),
            ("OTP 完美保密条件：真正随机、与消息等长、只用一次、密钥均匀分布且保密。", "One-time pad needs a truly random, message-length, one-use, secret key."),
            ("ECB 弱点：相同明文块产生相同密文块，泄露模式。", "ECB leaks patterns because identical plaintext blocks encrypt to identical ciphertext blocks."),
            ("Hash 用途：完整性检测、密码存储；核心性质是 preimage resistance、second-preimage resistance、collision resistance。", "Hash functions support integrity checks and password storage; know preimage, second-preimage, and collision resistance."),
            ("密码存储：hash 优于 encryption，因为不能用密钥还原；salt 防 rainbow table，但不能阻止针对单个密码的字典攻击。", "Hashing is preferred over encryption for passwords; salt blocks precomputed rainbow tables but not per-user dictionary attacks."),
            ("MAC/签名：MAC 需要共享密钥；数字签名用私钥签名、公钥验证。", "MAC uses a shared secret; digital signatures sign with private key and verify with public key."),
            ("对称 vs 非对称：对称适合大文件；非对称解决密钥交换和身份认证问题。", "Symmetric crypto is efficient for large data; asymmetric crypto solves key exchange and authentication."),
            ("DH：在不安全信道上协商共享秘密，但本身需要认证防 MITM。", "Diffie-Hellman establishes a shared secret over an insecure channel but needs authentication against MITM."),
        ]),
        ("认证、PKI、TLS / Authentication, PKI, TLS", [
            ("Root CA 私钥泄露：攻击者可伪造任意站点证书，但不能自动解密所有历史通信。", "If a root CA private key is stolen, attackers can sign fake certificates, but do not automatically decrypt all past traffic."),
            ("PKI 解决公钥真实性问题；浏览器预装 Root CA 公钥。", "PKI authenticates public keys; browsers/devices ship with root CA public keys."),
            ("TLS 作用：协商算法、认证服务器证书、建立对称加密通信、保护机密性和完整性。", "TLS negotiates algorithms, authenticates certificates, establishes symmetric keys, and protects confidentiality and integrity."),
            ("考试陷阱：TLS 不是专门负责发送账号密码；账号密码只是被 TLS 保护的应用层数据。", "Trap: TLS does not exist to send account/password; it protects application-layer data."),
        ]),
        ("访问控制 / Access Control", [
            ("Access Control 保护 CIA 三方面：机密性、完整性、可用性。", "Access control protects confidentiality, integrity, and availability."),
            ("Access Control Matrix 大系统低效：大量空单元，扩展性差。", "Access control matrices are inefficient at scale because they are sparse."),
            ("DAC：资源所有者可授权；MAC：管理员/系统强制策略，用户不能覆盖；RBAC：按角色授权。", "DAC lets owners grant rights; MAC enforces system policy; RBAC grants rights by role."),
            ("Bell-LaPadula 偏机密性：no read up, no write down；Biba 偏完整性：no read down, no write up。", "Bell-LaPadula protects confidentiality; Biba protects integrity."),
            ("ACL vs Capability：ACL 以对象列用户权限；Capability 以主体持有可访问对象能力。", "ACLs attach rights to objects; capabilities attach rights/tokens to subjects."),
        ]),
        ("浏览器与 Web 安全 / Browser and Web Security", [
            ("Origin = scheme + host + port；路径不同不影响 origin。", "Origin equals scheme + host + port; path does not define origin."),
            ("SOP：网页可以向任意网站发送请求，但只能读取同源响应。", "SOP allows sending requests broadly but reading only same-origin responses."),
            ("第三方脚本执行 origin：脚本嵌入在哪个页面，就以该页面 origin 执行。", "A third-party script executes under the embedding page's origin."),
            ("CSRF：利用用户已登录 cookie，让浏览器发出伪造请求；服务器只看到请求和 cookie，无法自然知道点击来自攻击站点。", "CSRF abuses authenticated cookies; the server sees only request plus cookie and cannot naturally infer the click origin."),
            ("CSRF 防御：CSRF token、SameSite cookie、检查 Origin/Referer、关键操作重新认证。", "CSRF defenses: tokens, SameSite cookies, Origin/Referer checks, re-authentication."),
            ("XSS vs SQLi：二者都来自数据/代码边界失败；XSS 在浏览器执行脚本，SQLi 在数据库解释 SQL。", "Both XSS and SQLi confuse data and code; XSS executes in browsers, SQLi in databases."),
        ]),
        ("网络安全 / Network Security", [
            ("DNS 原始设计缺少加密和认证，容易被欺骗；响应必须匹配源/目的 IP、端口、Query ID。", "Standard DNS lacks encryption/authentication; forged responses must match IPs, ports, and query ID."),
            ("Bailiwick check：权威服务器只能提供其管辖域内记录，不能为无关域背书。", "Bailiwick checking limits records to the server's authoritative domain."),
            ("Kaminsky Attack：查询随机不存在子域以绕过缓存，重复制造竞态窗口。", "Kaminsky attack queries random nonexistent subdomains to bypass cache and create repeated races."),
            ("DNSSEC：签名 DNS 响应用于认证，但不解决 DNS-based DDoS。", "DNSSEC authenticates DNS data but does not solve DNS-based DDoS."),
            ("BGP：标准 BGP 缺少认证，易 prefix hijacking；RPKI 验证 IP 前缀/AS 授权。", "Standard BGP lacks authentication and is vulnerable to prefix hijacking; RPKI validates prefix/AS authorization."),
            ("DNS Amplification：攻击者伪造 victim 源 IP 给 DNS server，服务器把更大响应打向 victim；ISP 可用 ingress filtering/BCP38。", "DNS amplification spoofs the victim as source; ingress filtering/BCP38 prevents spoofed-source packets."),
            ("Worm 活动可通过 unused IP background radiation 观察。", "Worm activity can be measured by background radiation to unused IP space."),
        ]),
        ("IDS 与防火墙 / IDS and Firewalls", [
            ("Signature/misuse IDS：已知攻击准确、误报低；缺点是漏掉未知攻击且需更新签名。", "Signature/misuse IDS detects known attacks with high precision but misses unknown attacks and needs updates."),
            ("Anomaly IDS：可能发现未知攻击；缺点是误报较高，需要基线/训练。", "Anomaly IDS may detect unknown attacks but usually has more false positives and needs baselines/training."),
            ("Stateless firewall 单包判断；Stateful firewall 根据连接状态判断。", "Stateless firewalls inspect packets in isolation; stateful firewalls track connection context."),
            ("Default-deny 更安全但可能影响合法用户；严格信任内网会放过已沦陷内部主机。", "Default-deny is safer but may hurt usability; trusting all inside nodes fails when an internal host is compromised."),
            ("Zero Trust：不因设备在内网就默认可信，要持续验证。", "Zero Trust does not trust a device merely because it is inside the network; it validates continuously."),
        ]),
        ("系统/软件安全补漏 / System and Software Security Patch", [
            ("Shellcode：攻击者注入并执行的机器码；常见概念包括 NOP sled、返回地址覆盖、权限/系统调用。", "Shellcode is injected machine code; know NOP sleds, return-address overwrite, privileges, and syscalls."),
            ("Memory errors：缓冲区溢出、越界写、use-after-free 等会破坏控制流或数据完整性。", "Memory errors such as buffer overflow, out-of-bounds writes, and use-after-free can corrupt control flow or data."),
            ("防御：ASLR、stack canary、DEP/NX、sandbox、最小权限。", "Defenses: ASLR, stack canaries, DEP/NX, sandboxing, least privilege."),
            ("Race condition/TOCTOU：检查和使用之间状态改变；防御靠原子操作、锁、避免临时文件不安全用法。", "Race conditions/TOCTOU happen when state changes between check and use; use atomic operations and locks."),
        ]),
    ]
    for title_text, bullets in topics:
        add_h(doc, title_text, 2)
        for cn, en in bullets:
            add_bullet(doc, cn, en)

    add_h(doc, "4. 题型模板 / Question Templates", 1)
    add_table(doc, ["题型 / Type", "答题模板 / Answer Template", "常见考点 / Common Topics"], [
        ("判断题 T/F", "先找绝对词：always/only/all/never；再对照性质。", "Kerckhoff, SOP, DNSSEC, BGP, IDS, Zero Trust"),
        ("选择题 MCQ", "定义题直接背；场景题先识别目标：机密性/完整性/认证/可用性。", "Crypto choice, PKI/TLS, access control model, firewall type"),
        ("短答题 Short answer", "按 Cause -> Attack/Impact -> Defense 三段写。", "ECB, CSRF, XSS vs SQLi, Kaminsky, DNS amplification"),
        ("计算题 Calculation", "写出搜索空间或指数差，再换算数量级。", "DES 56 vs 112, PIN/hash brute force, birthday bound"),
        ("设计题 Design", "先说明 threat model，再列机制，不要只写工具名。", "Password storage, TLS use, firewall policy, IDS rule"),
    ], widths=[Inches(1.35), Inches(3.1), Inches(2.3)])

    add_h(doc, "5. 必背术语 / Bilingual Glossary", 1)
    glossary = [
        ("Confidentiality", "机密性", "Only authorized parties can read data."),
        ("Integrity", "完整性", "Data cannot be modified without detection/authorization."),
        ("Availability", "可用性", "Systems remain accessible when needed."),
        ("Kerckhoff's Principle", "柯克霍夫原则", "Security should not depend on algorithm secrecy."),
        ("One-wayness / Preimage resistance", "单向性/原像抗性", "Given h(x), hard to recover x."),
        ("Collision resistance", "碰撞抗性", "Hard to find two different inputs with same hash."),
        ("MAC", "消息认证码", "Shared-key integrity/authentication tag."),
        ("Digital Signature", "数字签名", "Private-key signature, public-key verification."),
        ("Certificate Authority", "证书颁发机构", "Trusted entity that signs certificates."),
        ("Same Origin Policy", "同源策略", "Restricts reading responses across origins."),
        ("CSRF", "跨站请求伪造", "Forces authenticated browser to send unwanted request."),
        ("XSS", "跨站脚本", "Injects script executed in user's browser."),
        ("SQL Injection", "SQL 注入", "Injects SQL interpreted by database."),
        ("Bailiwick Check", "管辖范围检查", "Rejects unrelated DNS records from a nameserver."),
        ("Prefix Hijacking", "前缀劫持", "False BGP route announcement for IP prefix."),
        ("Ingress Filtering", "入口过滤", "Drops packets with spoofed source IPs."),
        ("Misuse Detection", "误用检测/签名检测", "Detects known attack patterns."),
        ("Anomaly Detection", "异常检测", "Detects deviation from normal behavior."),
        ("Stateful Firewall", "有状态防火墙", "Tracks connection/session state."),
        ("TOCTOU", "检查到使用时差漏洞", "Race between checking a condition and using it."),
    ]
    add_table(doc, ["English", "中文", "Meaning"], glossary, widths=[Inches(1.65), Inches(1.55), Inches(3.55)])

    add_h(doc, "6. Practice / Quiz 答案索引", 1)
    add_para(doc, "Use this section for fast checking only. Real exam preparation requires explaining the reasoning. 本节只用于快速核对；真正备考必须能解释原因。", size=9.5, bold=True)
    add_table(doc, ["Source", "Answer Key / Main Key", "Note"], extract_answer_keys(data), widths=[Inches(2.55), Inches(2.2), Inches(2.0)])

    add_h(doc, "7. Practice 短答必会点 / Short-Answer Must-Knows", 1)
    short_points = [
        ("ECB weakness", "Identical plaintext blocks -> identical ciphertext blocks -> pattern leakage.", "相同明文块产生相同密文块，泄露统计/图案信息。"),
        ("Hash alone for integrity", "Plain hash sent with plaintext is not enough because attacker can modify both message and hash.", "明文+普通 hash 不能防篡改，因为攻击者可同时改消息和 hash。"),
        ("Why asymmetric crypto is used", "It solves key distribution/authentication even though it is slower and uses longer keys.", "非对称虽然慢、密钥长，但解决密钥分发和认证。"),
        ("Password salt failure example", "If stored value can be transformed back to hash(password), offline dictionary attack still works.", "若攻击者能从存储值还原 hash(password)，离线字典攻击仍成立。"),
        ("CSRF explanation", "Server receives same authenticated request/cookie, so it cannot know whether click came from bank or attacker site unless extra defenses exist.", "服务器只看到请求和 cookie，缺少额外防御时无法判断点击来源。"),
        ("XSS vs SQLi", "Both fail to separate data from code; XSS affects browser/client, SQLi affects database/server.", "二者都是数据/代码边界失败；XSS 在客户端，SQLi 在服务端数据库。"),
        ("Kaminsky attack", "Random nonexistent subdomains bypass cache and repeatedly trigger races against authoritative DNS replies.", "随机不存在子域绕过缓存，反复制造伪造响应竞态。"),
        ("DNS amplification", "Spoof victim source IP in DNS query; DNS server sends larger responses to victim; ISP ingress filtering blocks spoofing.", "伪造受害者源 IP 发 DNS 请求，服务器把放大响应打向受害者；入口过滤可防。"),
        ("Signature IDS", "High precision for known attacks, weak for unknown attacks and requires signature updates.", "已知攻击准确，未知攻击弱，需更新签名。"),
        ("Stateful vs stateless firewall", "Stateless sees one packet; stateful sees packet in connection context.", "无状态看单包；有状态看连接上下文。"),
    ]
    add_table(doc, ["Concept", "EN answer", "中文答案"], short_points, widths=[Inches(1.45), Inches(2.65), Inches(2.65)])

    add_h(doc, "8. 资料清单与缺口 / Material Inventory and Gaps", 1)
    ppt_titles = data.get("ppt_titles", [])
    add_bullet(doc, f"Slides/PPT present: {len(ppt_titles)} files, CSE331-01 through CSE331-20.", f"Slides/PPT present: {len(ppt_titles)} files, CSE331-01 through CSE331-20.")
    add_bullet(doc, "Practice present: Crypto, Access Control/Web Security, Network Security.", "Practice sets present: Crypto, Access Control/Web Security, Network Security.")
    add_bullet(doc, "Quiz present: Quiz 1-3, both Class 1 and Class 2 solutions.", "Quizzes present: Quiz 1-3, both class versions.")
    add_bullet(doc, "Homework present: MD5 collision, CSRF, DNS, Firewall, Shellcode, Buffer Overflow Server, Race Condition.", "Homework topics present: MD5 collision, CSRF, DNS, firewall, shellcode, buffer overflow server, race condition.")
    add_bullet(doc, "Missing or not found locally: Quiz 4-6 and later-topic slides beyond CSE331-20 if the instructor posted them separately.", "Missing or not found locally: Quiz 4-6 and any slides beyond CSE331-20.")

    doc.add_page_break()
    add_h(doc, "9. 考前最终检查 / Final Checklist", 1)
    for cn, en in [
        ("我能解释每个 Quiz 错项为什么错。", "I can explain why each wrong quiz option is wrong."),
        ("我能默写 ECB、CSRF、Kaminsky、DNS amplification、stateful firewall 的标准短答。", "I can write standard short answers for ECB, CSRF, Kaminsky, DNS amplification, and stateful firewalls."),
        ("我能区分 hash/MAC/signature/encryption/TLS/PKI 的作用。", "I can distinguish hash, MAC, signature, encryption, TLS, and PKI."),
        ("我能判断 SOP 中两个 URL 是否同源。", "I can decide whether two URLs have the same origin."),
        ("我能写出 DNS 伪造响应时 IP/端口/Query ID 如何匹配。", "I can fill matching IPs, ports, and Query ID for a forged DNS response."),
        ("我已扫过 shellcode、sandbox、memory errors、race condition 的定义和防御。", "I have skimmed definitions and defenses for shellcode, sandbox, memory errors, and race condition."),
    ]:
        add_bullet(doc, cn, en)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
